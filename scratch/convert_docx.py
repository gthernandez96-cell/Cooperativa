import re
import os
import sys

# Ensure python-docx is installed
try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed. Installing it now...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_formatted_run(paragraph, text, is_bold=False):
    """Helper to add runs with inline bold markdown parsing."""
    # Split text by bold markers **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def add_paragraph_with_formatting(doc, text, style=None, is_bullet=False):
    """Adds a paragraph with basic inline formatting."""
    if is_bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph(style=style)
    
    # Adjust spacing
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    
    add_formatted_run(p, text)
    return p

def main():
    # Paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(script_dir)
    
    # Locate the md document
    md_path = "/Users/gustavohernandez/.gemini/antigravity-ide/brain/508652b4-b057-42ea-bf04-29a8ca1fc029/documentacion_sistema.md"
    docx_output_path = os.path.join(project_dir, "CoopAhorro_Documentacion.docx")
    
    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found at {md_path}")
        sys.exit(1)
        
    print(f"Reading documentation from: {md_path}")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles Setup
    styles = doc.styles
    
    # Normal / Paragraph Font
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(51, 51, 51) # Charcoal
    
    # Primary theme colors
    primary_color = RGBColor(27, 54, 93)     # Navy Blue
    secondary_color = RGBColor(65, 105, 225) # Royal Blue
    
    # State tracking
    in_code_block = False
    in_table = False
    table_rows = []
    code_block_content = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # ── 1. Code Block (Mermaid, etc.) ──────────────────────
        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                # Process code block
                if code_block_content:
                    # If it's a mermaid graph, we can summarize it in a blockquote-like style
                    content_text = "\n".join(code_block_content)
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    
                    # Create border/background visual container
                    run = p.add_run("[Diagrama / Flujo estructurado]\n")
                    run.bold = True
                    run.font.color.rgb = secondary_color
                    
                    # Replace diagram code with cleaner representation
                    clean_content = ""
                    if "graph TD" in content_text or "flowchart" in content_text:
                        clean_content = "Representación del Flujo de Datos del Sistema:\n"
                        clean_content += "• Usuario/Cliente -> app.py (Flask Server)\n"
                        clean_content += "• app.py -> Middleware de Seguridad (CSRF & Rate Limiter)\n"
                        clean_content += "• app.py -> Módulos de Negocio (Socios, Ahorro, Préstamos, Configuración, etc.)\n"
                        clean_content += "• Módulos -> utils (db.py / helpers.py)\n"
                        clean_content += "• utils -> SQLite / PostgreSQL Database"
                    elif "stateDiagram" in content_text:
                        clean_content = "Representación de los Estados de un Préstamo:\n"
                        clean_content += "1. Registro de Solicitud -> Pendiente de aprobación\n"
                        clean_content += "2. Pendiente -> Aprobado por Comité (o marcado como No Procede si no califica)\n"
                        clean_content += "3. Aprobado -> Activo (Tras desembolso de fondos)\n"
                        clean_content += "4. Activo -> Al Día (Pagos regulares) / En Mora (Cuotas vencidas)\n"
                        clean_content += "5. En Mora -> Cobranza (Gestión telefónica y compromisos)\n"
                        clean_content += "6. Cobranza -> Legal (Tras 90+ días de retraso)\n"
                        clean_content += "7. Activo / Al Día -> Cancelado (Liquidación total y emisión de finiquito)"
                    else:
                        # Standard code block representation
                        clean_content = content_text
                        
                    p.add_run(clean_content)
                    p.runs[-1].font.italic = True
                    p.runs[-1].font.size = Pt(10)
                code_block_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_block_content.append(line.rstrip('\n'))
            i += 1
            continue
            
        # ── 2. Tables ──────────────────────────────────────────
        if stripped.startswith("|"):
            in_table = True
            # Parse table row
            parts = [p.strip() for p in stripped.split("|")[1:-1]]
            # If it's a separator line (e.g., | :--- | :--- |), skip it
            if all(re.match(r'^[-:]+$', p) for p in parts):
                i += 1
                continue
            table_rows.append(parts)
            i += 1
            continue
        else:
            if in_table:
                # Flush the table into the document
                if table_rows:
                    cols_count = len(table_rows[0])
                    table = doc.add_table(rows=len(table_rows), cols=cols_count)
                    table.style = 'Table Grid'
                    table.autofit = True
                    
                    for r_idx, row_data in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(table.columns):
                                cell = table.cell(r_idx, c_idx)
                                cell.text = "" # Clear default
                                p = cell.paragraphs[0]
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after = Pt(2)
                                add_formatted_run(p, cell_text)
                                
                                # Styling table headers
                                if r_idx == 0:
                                    set_cell_background(cell, "1B365D") # Navy blue
                                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                                    p.runs[0].bold = True
                                else:
                                    # Alternating rows
                                    if r_idx % 2 == 0:
                                        set_cell_background(cell, "F2F5F8")
                                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                                
                table_rows = []
                in_table = False
                
        # ── 3. Headings ────────────────────────────────────────
        if stripped.startswith("#"):
            # Determine heading level
            lvl = len(re.match(r'^#+', stripped).group(0))
            title_text = stripped.lstrip("#").strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if lvl == 1:
                run = p.add_run(title_text)
                run.font.size = Pt(22)
                run.bold = True
                run.font.color.rgb = primary_color
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                # Add horizontal line underneath main title
                p_border = doc.add_paragraph()
                p_border.paragraph_format.space_after = Pt(6)
            elif lvl == 2:
                run = p.add_run(title_text)
                run.font.size = Pt(16)
                run.bold = True
                run.font.color.rgb = primary_color
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            elif lvl == 3:
                run = p.add_run(title_text)
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = secondary_color
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            else:
                run = p.add_run(title_text)
                run.font.size = Pt(11)
                run.bold = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                
            i += 1
            continue
            
        # ── 4. Bullet lists ────────────────────────────────────
        if stripped.startswith("* ") or stripped.startswith("- "):
            list_text = stripped[2:].strip()
            add_paragraph_with_formatting(doc, list_text, is_bullet=True)
            i += 1
            continue
            
        # ── 5. Separators ──────────────────────────────────────
        if stripped == "---":
            # Just add a paragraph space
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
            
        # ── 6. Regular Paragraphs ──────────────────────────────
        if stripped:
            add_paragraph_with_formatting(doc, stripped)
            
        i += 1

    # Save Document
    doc.save(docx_output_path)
    print(f"Successfully generated Word document at: {docx_output_path}")

if __name__ == "__main__":
    main()
